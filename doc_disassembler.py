import docx
from enum import Enum
import re
import logger

DEFAULT_IMPORTANCE = str(5)
POINT_SPLIT = '#'

SELECT_REG = r'(^|\s)[ABCDEFGH][\.\)）]'
GRAMA_SELECT_ANS_REG = r'([ABCD][^E-Za-z]*?){19}[ABCD]'
GRAMA_FISH_ANS_REG = r'([ABCDE][^F-Za-z]*?){7}[ABCDE]'

QUE_NAME = [r'.*?Listen.+?picture.*',
            r'.*?Listen to.+?choose.*',
            r'.*?Listen to.+?true.+?false.*',
            r'.*?Listen to.+?(complete.+?sentences|fill.+?blank).*',
            r'.*?(Phonetics|Vocabulary|Grammar).+?(Phonetics|Vocabulary|Grammar).+?(Phonetics|Vocabulary|Grammar).*',
            r'.*?Choose the best answer.*',
            r'.*?Complete.+?(word|phrase).+?box.*',
            r'.*?Complete.+?(proper|suitable).+?forms.*',
            r'.*?sentences as required.*',
            r'.*?Reading.+?Writing.*']


class QUE_TYPE(Enum):
    LISTEN_PIC = '听力-选图'
    LISTEN_SELE = '听力-选择'
    LISTEN_TF = '听力-TF'
    LISTEN_FILL = '听力-填空'
    GRAMA_SELE = '语法-选择'
    GRAMA_FISH = '语法-钓鱼'
    GRAMA_ALTER_WORD = '语法-改词'
    GRAMA_ALTER_SENT = '语法-改句'
    READ_SELE = '阅读-选择'
    READ_PICK_WORD = '阅读-完型'
    READ_INIT = '阅读-首字母'
    READ_ANS = '阅读-回答问题'
    WRITE = '作文'


def type2num(type):
    if type == QUE_TYPE.LISTEN_PIC:
        return 6
    elif type == QUE_TYPE.LISTEN_SELE:
        return 8
    elif type == QUE_TYPE.LISTEN_TF:
        return 6
    elif type == QUE_TYPE.LISTEN_FILL:
        return 5
    elif type == QUE_TYPE.GRAMA_SELE:
        return 20
    elif type == QUE_TYPE.GRAMA_FISH:
        return 8
    elif type == QUE_TYPE.GRAMA_ALTER_WORD:
        return 8
    elif type == QUE_TYPE.GRAMA_ALTER_SENT:
        return 7
    elif type == QUE_TYPE.READ_SELE:
        return 6
    elif type == QUE_TYPE.READ_PICK_WORD:
        return 6
    elif type == QUE_TYPE.READ_INIT:
        return 7
    elif type == QUE_TYPE.READ_ANS:
        return 6
    elif type == QUE_TYPE.WRITE:
        return 1


def devide_para(para):
    text = []
    point = []
    imp = []
    for r in para.runs:
        if str(r.font.color.rgb) == 'FF0000':
            point.append(r.text)
        elif str(r.font.color.rgb) == '00B0F0':
            imp.append(r.text)
        elif r.underline:
            text.append('<_>' + r.text + '</_>')
        else:
            text.append(r.text)
    text = ''.join(text)
    point = ''.join(point)
    imp = ''.join(imp)
    text = re.sub(r'\(\s\)', '', text)
    text = re.sub(r'（\s）', '', text)
    return text.strip(), point.strip(), imp.strip()


def pick_para_out(docu, start, left, right):
    res = []
    flag = 'init'
    i = start
    while i < len(docu.paragraphs):
        para = docu.paragraphs[i]
        i += 1
        if len(para.text.strip()) == 0:
            continue
        tmp_text = re.sub(r'\s', ' ', para.text.strip())
        if flag == 'init':
            if re.match(QUE_NAME[left], tmp_text, re.I):
                flag = 'in'
            else:
                continue
        elif flag == 'in':
            if re.match(QUE_NAME[right], tmp_text, re.I):
                flag = 'out'
                break
            else:
                res.append(para)
    if flag != 'out':
        logger.log('未找到合法题型名称', QUE_NAME[left] + '\n' + QUE_NAME[right])
        raise RuntimeError('未找到合法题型名称')
    else:
        return res, i - 1


def get_text_type(para):
    for run in para.runs:
        if run.underline and len(re.sub(r'[\d]+', '', run.text).strip()) == 0:
            return 'normal'

    text = para.text
    bold_len = 0
    for run in para.runs:
        if run.bold:
            bold_len += len(run.text)
    if bold_len / len(text) > 0.8 and '分' in text:
        return 'bold'

    text = re.sub(r'[\d]+\.', '', text.strip())
    if re.match(SELECT_REG, text):
        sele_num = re.findall(SELECT_REG, text)
        try_sele = re.split(SELECT_REG, text.strip())
        try_sele = list(filter(lambda item: len(item.strip()) > 0, try_sele))
        if len(try_sele) >= len(sele_num):
            return 'select'

    return 'normal'


def generate_select(s):
    tmp_try_end = s.strip()
    tmp_try_end = re.split(SELECT_REG, tmp_try_end)
    tmp_try_end = list(map(lambda item: re.sub(r'\s\.$', '.', item.strip()).strip(), tmp_try_end))
    tmp_try_end = list(filter(lambda item: len(item) > 0, tmp_try_end))

    return tmp_try_end


def read_with_one_doc(a):
    main = docx.Document(a)

    # 读取试卷

    lis_sele, end = pick_para_out(main, 0, 1, 2)
    lis_tf, end = pick_para_out(main, end, 2, 3)
    lis_fill, end = pick_para_out(main, end, 3, 4)
    gra_sele, end = pick_para_out(main, end, 5, 6)
    gra_alt_word, end = pick_para_out(main, end, 7, 8)
    gra_alt_sent, paper_end = pick_para_out(main, end, 8, 9)

    # lis_sele_arr = []
    # tmp_q = ''
    # for para in lis_sele:
    #     text, point, imp = devide_para(para)
    #     text = re.sub(r'^\d+?\.', '', text)
    #     tmp_q += text
    #     tmp_sele = generate_select(tmp_q)
    #     if len(tmp_sele) == 4:
    #         lis_sele_arr.append(tmp_sele)
    #         tmp_q = ''
    #     elif len(tmp_sele) > 4:
    #         raise RuntimeError('选项数大于4')
    #     else:
    #         tmp_q += '\n'
    # if len(lis_sele_arr) != type2num(QUE_TYPE.LISTEN_SELE):
    #     raise RuntimeError(QUE_TYPE.LISTEN_SELE.value+'题数错误')
    # print(lis_sele_arr)
    #
    # lis_tf_arr = []
    # for para in lis_tf:
    #     text, point, imp = devide_para(para)
    #     text = re.sub(r'^\d+?\.', '', text).strip()
    #     lis_tf_arr.append(text)
    # if len(lis_tf_arr) != type2num(QUE_TYPE.LISTEN_TF):
    #     raise RuntimeError(QUE_TYPE.LISTEN_TF.value+'题数错误')
    # print(lis_tf_arr)
    #
    # lis_fill_arr = []
    # for para in lis_fill:
    #     text, point, imp = devide_para(para)
    #     text = re.sub(r'^\d+?\.', '', text).strip()
    #     lis_fill_arr.append(text)
    # if len(lis_fill_arr) != type2num(QUE_TYPE.LISTEN_FILL):
    #     raise RuntimeError(QUE_TYPE.LISTEN_FILL.value+'题数错误')
    # print(lis_fill_arr)

    gra_sele_arr = []
    gra_sele_point = []
    gra_sele_imp = []
    tmp_q = ['', '']
    tmp_p = []
    tmp_i = DEFAULT_IMPORTANCE
    for para in gra_sele:
        text, point, imp = devide_para(para)
        text = re.sub(r'^\d+?\.', '', text)
        if get_text_type(para) == 'select':
            tmp_q[1] += text
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
            tmp_sele = generate_select(tmp_q[1])
            if len(tmp_sele) == 4:
                tmp_q[1] = tmp_sele[:]
                tmp_q[0] = tmp_q[0].strip()
                gra_sele_arr.append(tmp_q[:])
                gra_sele_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
                gra_sele_imp.append(tmp_i)
                tmp_q = ['', '']
                tmp_p = []
                tmp_i = DEFAULT_IMPORTANCE
            elif len(tmp_sele) > 4:
                logger.log('选项数大于4',
                           '\n'.join([item.text for item in gra_sele]) + '\n' + '#'.join(tmp_sele) + '\n' + text)
                raise RuntimeError('选项数大于4')
            else:
                tmp_q[1] += '\n'
        else:
            tmp_q[0] += '\n' + text
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
    if len(gra_sele_arr) != type2num(QUE_TYPE.GRAMA_SELE):
        logger.log(QUE_TYPE.GRAMA_SELE.value + '题数错误',
                   '\n'.join([item.text for item in gra_sele]) + '\n' + '\n'.join(gra_sele_arr))
        raise RuntimeError(QUE_TYPE.GRAMA_SELE.value + '题数错误')

    gra_fish_path = a

    gra_alt_word_arr = []
    gra_alt_word_point = []
    gra_alt_word_imp = []
    tmp_q = ''
    tmp_p = []
    tmp_i = DEFAULT_IMPORTANCE
    for para in gra_alt_word:
        text, point, imp = devide_para(para)
        text = re.sub(r'^\d+?\.', '', text).strip()
        tmp_q += text
        tmp_p.append(point.strip())
        tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
        if tmp_q[-1] == ')' or tmp_q[-1] == '）':
            gra_alt_word_arr.append(tmp_q)
            gra_alt_word_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
            gra_alt_word_imp.append(tmp_i)
            tmp_q = ''
            tmp_p = []
            tmp_i = DEFAULT_IMPORTANCE
        else:
            tmp_q += '\n'
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
    if len(gra_alt_word_arr) != type2num(QUE_TYPE.GRAMA_ALTER_WORD):
        logger.log(QUE_TYPE.GRAMA_ALTER_WORD.value + '题数错误',
                   '\n'.join([item.text for item in gra_alt_word]) + '\n' + '\n'.join(gra_alt_word_arr))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_WORD.value + '题数错误')

    gra_alt_sent_arr = []
    gra_alt_sent_point = []
    gra_alt_sent_imp = []
    tmp_q = ['', '']
    tmp_p = []
    tmp_i = DEFAULT_IMPORTANCE
    state = 'question'
    i = 0
    while i < len(gra_alt_sent):
        para = gra_alt_sent[i]
        i += 1
        text, point, imp = devide_para(para)
        text = re.sub(r'^\d+?\.', '', text).strip()
        if state == 'question':
            tmp_q[0] += '\n' + text
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
            if text[-1] == ')' or text[-1] == '）':
                state = 'answer'
        elif state == 'answer':
            if text[-1] == ')' or text[-1] == '）':
                tmp_q[0] = tmp_q[0].strip()
                tmp_q[1] = tmp_q[1].strip()
                gra_alt_sent_arr.append(tmp_q[:])
                gra_alt_sent_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
                gra_alt_sent_imp.append(tmp_i)
                tmp_q = ['', '']
                tmp_p = []
                tmp_i = DEFAULT_IMPORTANCE
                state = 'question'
                i -= 1
            else:
                tmp_q[1] += '\n' + text
                tmp_p.append(point.strip())
                tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
    tmp_q[0] = tmp_q[0].strip()
    tmp_q[1] = tmp_q[1].strip()
    gra_alt_sent_arr.append(tmp_q[:])
    gra_alt_sent_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
    gra_alt_sent_imp.append(tmp_i)
    if len(gra_alt_sent_arr) != type2num(QUE_TYPE.GRAMA_ALTER_SENT):
        logger.log(QUE_TYPE.GRAMA_ALTER_SENT.value + '题数错误',
                   '\n'.join([item.text for item in gra_alt_sent]) + '\n#########\n' + '\n'.join(
                       list(map(lambda item: '\n'.join(item), gra_alt_sent_arr))))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_SENT.value + '题数错误')

    # 读取答案

    # lis_pic, end = pick_para_out(main, paper_end, 0, 1)
    # lis_sele, end = pick_para_out(main, paper_end, 1, 2)
    # lis_tf, end = pick_para_out(main, paper_end, 2, 3)
    # for p in lis_tf:
    #     print(p.text)
    ans_para = main.paragraphs[paper_end:]
    ans_para = list(map(lambda item: item.text, ans_para))
    ans_para = '\n'.join(ans_para)

    gra_sele_ans = re.search(GRAMA_SELECT_ANS_REG, ans_para)
    if not gra_sele_ans:
        logger.log(QUE_TYPE.GRAMA_SELE.value + '答案数错误', ans_para)
        raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_SELE.value + '答案')
    end = gra_sele_ans.span()[1]
    gra_sele_ans = gra_sele_ans.group(0)
    gra_sele_ans = re.sub(r'[^A-D]', '', gra_sele_ans)
    if len(gra_sele_ans) != type2num(QUE_TYPE.GRAMA_SELE):
        logger.log(QUE_TYPE.GRAMA_SELE.value + '答案数错误', ans_para + '\n' + gra_sele_ans)
        raise RuntimeError(QUE_TYPE.GRAMA_SELE.value + '答案数错误')

    ans_para = ans_para[end:]

    gra_fish_ans = re.search(GRAMA_FISH_ANS_REG, ans_para)
    if not gra_fish_ans:
        logger.log(QUE_TYPE.GRAMA_FISH.value + '答案数错误', ans_para)
        raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_FISH.value + '答案')
    end = gra_fish_ans.span()[1]
    gra_fish_ans = gra_fish_ans.group(0)
    gra_fish_ans = re.sub(r'[^A-E]', '', gra_fish_ans)
    if len(gra_fish_ans) != type2num(QUE_TYPE.GRAMA_FISH):
        logger.log(QUE_TYPE.GRAMA_FISH.value + '答案数错误', ans_para + '\n' + gra_fish_ans)
        raise RuntimeError(QUE_TYPE.GRAMA_FISH.value + '答案数错误')

    ans_para = ans_para[end:]

    gra_alt_word_ans = []
    start = 54
    for i in range(start, start + type2num(QUE_TYPE.GRAMA_ALTER_WORD)):
        if i == start + type2num(QUE_TYPE.GRAMA_ALTER_WORD) - 1:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\n)', ans_para)
        else:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\s*?)' + str(i + 1), ans_para)
        if not tmp_ans:
            logger.log('未能匹配出' + QUE_TYPE.GRAMA_ALTER_WORD.value + '答案', ans_para + '\n' + str(i))
            raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_ALTER_WORD.value + '答案')
        tmp_ans = tmp_ans.group(1)
        tmp_ans = re.sub('^[^a-zA-Z]+', '', tmp_ans).strip()
        gra_alt_word_ans.append(tmp_ans)
    if len(gra_alt_word_ans) != type2num(QUE_TYPE.GRAMA_ALTER_WORD):
        logger.log(QUE_TYPE.GRAMA_ALTER_WORD.value + '答案数错误', ans_para + '\n' + '#'.join(gra_alt_word_ans))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_WORD.value + '答案数错误')

    gra_alt_sent_ans = []
    start = 62
    for i in range(start, start + type2num(QUE_TYPE.GRAMA_ALTER_SENT)):
        if i == start + type2num(QUE_TYPE.GRAMA_ALTER_SENT) - 1:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\n)', ans_para)
        else:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\s*?)' + str(i + 1), ans_para)
        if not tmp_ans:
            logger.log('未能匹配出' + QUE_TYPE.GRAMA_ALTER_SENT.value + '答案', ans_para + '\n' + str(i))
            raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_ALTER_SENT.value + '答案')
        tmp_ans = tmp_ans.group(1)
        tmp_ans = re.sub('^[^a-zA-Z]+', '', tmp_ans).strip()
        gra_alt_sent_ans.append(tmp_ans)
    if len(gra_alt_sent_ans) != type2num(QUE_TYPE.GRAMA_ALTER_SENT):
        logger.log(QUE_TYPE.GRAMA_ALTER_SENT.value + '答案数错误', ans_para + '\n' + '#'.join(gra_alt_sent_ans))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_SENT.value + '答案数错误')

    return [{
        'question': gra_sele_arr,
        'point': gra_sele_point,
        'importance': gra_sele_imp,
        'answer': gra_sele_ans
    }, {
        'question': gra_fish_path,
        'point': None,
        'importance': None,
        'answer': gra_fish_ans
    }, {
        'question': gra_alt_word_arr,
        'point': gra_alt_word_point,
        'importance': gra_alt_word_imp,
        'answer': gra_alt_word_ans
    }, {
        'question': gra_alt_sent_arr,
        'point': gra_alt_sent_point,
        'importance': gra_alt_sent_imp,
        'answer': gra_alt_sent_ans
    }]


def read_with_two_doc(a, b):
    main = docx.Document(a)
    answer = docx.Document(b)

    # 读取试卷

    lis_sele, end = pick_para_out(main, 0, 1, 2)
    lis_tf, end = pick_para_out(main, end, 2, 3)
    lis_fill, end = pick_para_out(main, end, 3, 4)
    gra_sele, end = pick_para_out(main, end, 5, 6)
    gra_alt_word, end = pick_para_out(main, end, 7, 8)
    gra_alt_sent, _ = pick_para_out(main, end, 8, 9)

    # lis_sele_arr = []
    # tmp_q = ''
    # for para in lis_sele:
    #     text, point, imp = devide_para(para)
    #     text = re.sub(r'^\d+?\.', '', text)
    #     tmp_q += text
    #     tmp_sele = generate_select(tmp_q)
    #     if len(tmp_sele) == 4:
    #         lis_sele_arr.append(tmp_sele)
    #         tmp_q = ''
    #     elif len(tmp_sele) > 4:
    #         raise RuntimeError('选项数大于4')
    #     else:
    #         tmp_q += '\n'
    # if len(lis_sele_arr) != type2num(QUE_TYPE.LISTEN_SELE):
    #     raise RuntimeError(QUE_TYPE.LISTEN_SELE.value+'题数错误')
    # print(lis_sele_arr)
    #
    # lis_tf_arr = []
    # for para in lis_tf:
    #     text, point, imp = devide_para(para)
    #     text = re.sub(r'^\d+?\.', '', text).strip()
    #     lis_tf_arr.append(text)
    # if len(lis_tf_arr) != type2num(QUE_TYPE.LISTEN_TF):
    #     raise RuntimeError(QUE_TYPE.LISTEN_TF.value+'题数错误')
    # print(lis_tf_arr)
    #
    # lis_fill_arr = []
    # for para in lis_fill:
    #     text, point, imp = devide_para(para)
    #     text = re.sub(r'^\d+?\.', '', text).strip()
    #     lis_fill_arr.append(text)
    # if len(lis_fill_arr) != type2num(QUE_TYPE.LISTEN_FILL):
    #     raise RuntimeError(QUE_TYPE.LISTEN_FILL.value+'题数错误')
    # print(lis_fill_arr)

    gra_sele_arr = []
    gra_sele_point = []
    gra_sele_imp = []
    tmp_q = ['', '']
    tmp_p = []
    tmp_i = DEFAULT_IMPORTANCE
    for para in gra_sele:
        text, point, imp = devide_para(para)
        text = re.sub(r'^\d+?\.', '', text)
        if get_text_type(para) == 'select':
            tmp_q[1] += text
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
            tmp_sele = generate_select(tmp_q[1])
            if len(tmp_sele) == 4:
                tmp_q[1] = tmp_sele[:]
                tmp_q[0] = tmp_q[0].strip()
                gra_sele_arr.append(tmp_q[:])
                gra_sele_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
                gra_sele_imp.append(tmp_i)
                tmp_q = ['', '']
                tmp_p = []
                tmp_i = DEFAULT_IMPORTANCE
            elif len(tmp_sele) > 4:
                logger.log('选项数大于4',
                           '\n'.join([item.text for item in gra_sele]) + '\n' + '#'.join(tmp_sele) + '\n' + text)
                raise RuntimeError('选项数大于4')
            else:
                tmp_q[1] += '\n'
        else:
            tmp_q[0] += '\n' + text
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
    if len(gra_sele_arr) != type2num(QUE_TYPE.GRAMA_SELE):
        logger.log(QUE_TYPE.GRAMA_SELE.value + '题数错误',
                   '\n'.join([item.text for item in gra_sele]) + '\n' + '\n'.join(gra_sele_arr))
        raise RuntimeError(QUE_TYPE.GRAMA_SELE.value + '题数错误')

    gra_fish_path = a

    gra_alt_word_arr = []
    gra_alt_word_point = []
    gra_alt_word_imp = []
    tmp_q = ''
    tmp_p = []
    tmp_i = DEFAULT_IMPORTANCE
    for para in gra_alt_word:
        text, point, imp = devide_para(para)
        text = re.sub(r'^\d+?\.', '', text).strip()
        tmp_q += text
        tmp_p.append(point.strip())
        tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
        if tmp_q[-1] == ')' or tmp_q[-1] == '）':
            gra_alt_word_arr.append(tmp_q)
            gra_alt_word_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
            gra_alt_word_imp.append(tmp_i)
            tmp_q = ''
            tmp_p = []
            tmp_i = DEFAULT_IMPORTANCE
        else:
            tmp_q += '\n'
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
    if len(gra_alt_word_arr) != type2num(QUE_TYPE.GRAMA_ALTER_WORD):
        logger.log(QUE_TYPE.GRAMA_ALTER_WORD.value + '题数错误',
                   '\n'.join([item.text for item in gra_alt_word]) + '\n' + '\n'.join(gra_alt_word_arr))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_WORD.value + '题数错误')

    gra_alt_sent_arr = []
    gra_alt_sent_point = []
    gra_alt_sent_imp = []
    tmp_q = ['', '']
    tmp_p = []
    tmp_i = DEFAULT_IMPORTANCE
    state = 'question'
    i = 0
    while i < len(gra_alt_sent):
        para = gra_alt_sent[i]
        i += 1
        text, point, imp = devide_para(para)
        text = re.sub(r'^\d+?\.', '', text).strip()
        if state == 'question':
            tmp_q[0] += '\n' + text
            tmp_p.append(point.strip())
            tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
            if text[-1] == ')' or text[-1] == '）':
                state = 'answer'
        elif state == 'answer':
            if text[-1] == ')' or text[-1] == '）':
                tmp_q[0] = tmp_q[0].strip()
                tmp_q[1] = tmp_q[1].strip()
                gra_alt_sent_arr.append(tmp_q[:])
                gra_alt_sent_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
                gra_alt_sent_imp.append(tmp_i)
                tmp_q = ['', '']
                tmp_p = []
                tmp_i = DEFAULT_IMPORTANCE
                state = 'question'
                i -= 1
            else:
                tmp_q[1] += '\n' + text
                tmp_p.append(point.strip())
                tmp_i = imp.strip() if len(imp.strip()) > 0 else tmp_i
    tmp_q[0] = tmp_q[0].strip()
    tmp_q[1] = tmp_q[1].strip()
    gra_alt_sent_arr.append(tmp_q[:])
    gra_alt_sent_point.append(POINT_SPLIT.join(list(filter(lambda item: len(item) > 0, tmp_p))))
    gra_alt_sent_imp.append(tmp_i)
    if len(gra_alt_sent_arr) != type2num(QUE_TYPE.GRAMA_ALTER_SENT):
        logger.log(QUE_TYPE.GRAMA_ALTER_SENT.value + '题数错误',
                   '\n'.join([item.text for item in gra_alt_sent]) + '\n#########\n' + '\n'.join(
                       list(map(lambda item: '\n'.join(item), gra_alt_sent_arr))))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_SENT.value + '题数错误')

    # 读取答案

    # lis_pic, end = pick_para_out(main, paper_end, 0, 1)
    # lis_sele, end = pick_para_out(main, paper_end, 1, 2)
    # lis_tf, end = pick_para_out(main, paper_end, 2, 3)
    # for p in lis_tf:
    #     print(p.text)
    ans_para = answer.paragraphs
    ans_para = list(map(lambda item: item.text, ans_para))
    ans_para = '\n'.join(ans_para)

    gra_sele_ans = re.search(GRAMA_SELECT_ANS_REG, ans_para)
    if not gra_sele_ans:
        logger.log(QUE_TYPE.GRAMA_SELE.value + '答案数错误', ans_para)
        raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_SELE.value + '答案')
    end = gra_sele_ans.span()[1]
    gra_sele_ans = gra_sele_ans.group(0)
    gra_sele_ans = re.sub(r'[^A-D]', '', gra_sele_ans)
    if len(gra_sele_ans) != type2num(QUE_TYPE.GRAMA_SELE):
        logger.log(QUE_TYPE.GRAMA_SELE.value + '答案数错误', ans_para + '\n' + gra_sele_ans)
        raise RuntimeError(QUE_TYPE.GRAMA_SELE.value + '答案数错误')

    ans_para = ans_para[end:]

    gra_fish_ans = re.search(GRAMA_FISH_ANS_REG, ans_para)
    if not gra_fish_ans:
        logger.log(QUE_TYPE.GRAMA_FISH.value + '答案数错误', ans_para)
        raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_FISH.value + '答案')
    end = gra_fish_ans.span()[1]
    gra_fish_ans = gra_fish_ans.group(0)
    gra_fish_ans = re.sub(r'[^A-E]', '', gra_fish_ans)
    if len(gra_fish_ans) != type2num(QUE_TYPE.GRAMA_FISH):
        logger.log(QUE_TYPE.GRAMA_FISH.value + '答案数错误', ans_para + '\n' + gra_fish_ans)
        raise RuntimeError(QUE_TYPE.GRAMA_FISH.value + '答案数错误')

    ans_para = ans_para[end:]

    gra_alt_word_ans = []
    start = 54
    for i in range(start, start + type2num(QUE_TYPE.GRAMA_ALTER_WORD)):
        if i == start + type2num(QUE_TYPE.GRAMA_ALTER_WORD) - 1:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\n)', ans_para)
        else:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\s*?)' + str(i + 1), ans_para)
        if not tmp_ans:
            logger.log('未能匹配出' + QUE_TYPE.GRAMA_ALTER_WORD.value + '答案', ans_para + '\n' + str(i))
            raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_ALTER_WORD.value + '答案')
        tmp_ans = tmp_ans.group(1)
        tmp_ans = re.sub('^[^a-zA-Z]+', '', tmp_ans).strip()
        gra_alt_word_ans.append(tmp_ans)
    if len(gra_alt_word_ans) != type2num(QUE_TYPE.GRAMA_ALTER_WORD):
        logger.log(QUE_TYPE.GRAMA_ALTER_WORD.value + '答案数错误', ans_para + '\n' + '#'.join(gra_alt_word_ans))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_WORD.value + '答案数错误')

    gra_alt_sent_ans = []
    start = 62
    for i in range(start, start + type2num(QUE_TYPE.GRAMA_ALTER_SENT)):
        if i == start + type2num(QUE_TYPE.GRAMA_ALTER_SENT) - 1:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\n)', ans_para)
        else:
            tmp_ans = re.search(str(i) + r'(.+?[a-zA-z].*?\s*?)' + str(i + 1), ans_para)
        if not tmp_ans:
            logger.log('未能匹配出' + QUE_TYPE.GRAMA_ALTER_SENT.value + '答案', ans_para + '\n' + str(i))
            raise RuntimeError('未能匹配出' + QUE_TYPE.GRAMA_ALTER_SENT.value + '答案')
        tmp_ans = tmp_ans.group(1)
        tmp_ans = re.sub('^[^a-zA-Z]+', '', tmp_ans).strip()
        gra_alt_sent_ans.append(tmp_ans)
    if len(gra_alt_sent_ans) != type2num(QUE_TYPE.GRAMA_ALTER_SENT):
        logger.log(QUE_TYPE.GRAMA_ALTER_SENT.value + '答案数错误', ans_para + '\n' + '#'.join(gra_alt_sent_ans))
        raise RuntimeError(QUE_TYPE.GRAMA_ALTER_SENT.value + '答案数错误')

    return [{
        'question': gra_sele_arr,
        'point': gra_sele_point,
        'importance': gra_sele_imp,
        'answer': gra_sele_ans
    }, {
        'question': gra_fish_path,
        'point': None,
        'importance': None,
        'answer': gra_fish_ans
    }, {
        'question': gra_alt_word_arr,
        'point': gra_alt_word_point,
        'importance': gra_alt_word_imp,
        'answer': gra_alt_word_ans
    }, {
        'question': gra_alt_sent_arr,
        'point': gra_alt_sent_point,
        'importance': gra_alt_sent_imp,
        'answer': gra_alt_sent_ans
    }]

# tmp_path = 'E:/代码/exam_generater/2018初三英语各区一模/06-期末考试-201801-徐汇/'
# tmp_path = 'E:/代码/exam_generater/2017初三英语各区一模二模/2017二模电子稿/'
# read_with_one_doc(tmp_path + '17年崇明二模.docx', tmp_path + '06-期末考试-201801-徐汇-答案.docx',
#                   tmp_path + '06-期末考试-201801-徐汇-听力.docx')
