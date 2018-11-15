import os
import re
import json
import docx
from docx.oxml.ns import qn
from docx.shared import Pt
import math
import codecs
import random
import shutil
import doc_disassembler
import pandas as pd
from docx2xlsx import TXT_SPLIT, SELE_SPLIT, MAIN_XLS, copy_sheet

DST_PATH = os.getcwd() + '/结果/'
TAB_LEN = 4
SPACE_NUM = 6
SELE_LEN = {
    4: 16,
    2: 32
}
for k in SELE_LEN.keys():
    SELE_LEN[k] -= 3


def judge_dst_dir(dst_dir):
    if not os.path.exists(DST_PATH):
        return '', True
    dst_dir = dst_dir.strip()
    if len(dst_dir) == 0:
        return '请输入非空白文件名', False
    if '/' in dst_dir or '\\' in dst_dir:
        return '文件名不要出现正反斜杠', False
    else:
        dst_dir = DST_PATH + dst_dir + '/'
        if os.path.exists(dst_dir):
            return '此文件名已存在', False
        else:
            return '', True


def read_config(config_path):
    with codecs.open(config_path, 'r', 'utf-8') as f:
        res = json.load(f)
    return res


def is_all_done(has_searched, tmp_point, config, type):
    for p in tmp_point:
        if p in has_searched:
            continue
        else:
            if config['父子分类'][type].__contains__(p):
                return False
    return True


def get_point(config, type, i):
    i += 1
    if config['题目类型'][type].__contains__(str(i)):
        tmp_point = config['题目类型'][type][str(i)]
    elif config['题目类型'][type].__contains__('默认'):
        tmp_point = config['题目类型'][type]['默认']
    else:
        tmp_point = ''

    tmp_point = tmp_point.split(doc_disassembler.POINT_SPLIT)
    has_searched = set()
    while not is_all_done(has_searched, tmp_point, config, type):
        for p in tmp_point:
            if p in has_searched:
                continue
            else:
                has_searched.add(p)
                if config['父子分类'][type].__contains__(p):
                    tmp_point += config['父子分类'][type][p]
                else:
                    continue
    tmp_point = list(set(tmp_point))
    tmp_point = list(filter(lambda item: len(item.strip()) > 0, tmp_point))

    return tmp_point


def is_intersect(x, y):
    if len(list(set(x).intersection(set(y)))) == 0:
        return False
    else:
        return True


def special_select(df, point_arr):
    tmp = df.copy()
    tmp.drop(df.index, inplace=True)
    for i in range(len(df)):
        if is_intersect(df.iloc[i]['题目考点'].split(doc_disassembler.POINT_SPLIT), point_arr):
            tmp = tmp.append(pd.Series(df.iloc[i]), ignore_index=True)
    return tmp


def drop_picked(df, has_picked):
    drop_arr = []
    for i in range(len(df)):
        if df.iloc[i]['题目路径'] in has_picked:
            drop_arr.append(i)
    return df.drop(drop_arr)


def random_arr_with_weight(weight_arr):
    tmp_arr = []
    arr_sum = 0
    for x in weight_arr:
        arr_sum += x
        tmp_arr.append(arr_sum)
    weight_arr = [x / arr_sum for x in tmp_arr]
    rand = random.random()
    for i, v in enumerate(weight_arr):
        if rand < v:
            return i


def random_sele(sele_arr, ans):
    ans = ans.strip()
    tmp_arr = [(sele_arr[i], chr(i + ord('A'))) for i in range(len(sele_arr))]
    tmp_arr.sort(key=lambda item: random.random())
    new_ans = False
    for i, sele in enumerate(tmp_arr):
        if sele[1] == ans:
            new_ans = chr(i + ord('A'))
    tmp_arr = [tmp[0] for tmp in tmp_arr]
    return tmp_arr, new_ans


def path_to_question_and_answer(path, type):
    with codecs.open(path, 'r', 'utf-8') as f:
        file_info = f.read()
    file_info = file_info.split(TXT_SPLIT)
    file_info = list(map(lambda item: item.strip(), file_info))
    res = {}
    if type == doc_disassembler.QUE_TYPE.GRAMA_SELE.value:
        new_sele, new_ans = random_sele(list(map(lambda item: item.strip(), file_info[1].split(SELE_SPLIT))),
                                        file_info[2])
        res['question'] = [file_info[0], new_sele]
        res['answer'] = new_ans
    elif type == doc_disassembler.QUE_TYPE.GRAMA_FISH.value:
        res['answer'] = list(map(lambda item: item.strip(), file_info[0].split(SELE_SPLIT)))
    elif type == doc_disassembler.QUE_TYPE.GRAMA_ALTER_WORD.value:
        res['question'] = file_info[0]
        res['answer'] = file_info[1]
    elif type == doc_disassembler.QUE_TYPE.GRAMA_ALTER_SENT.value:
        res['question'] = [file_info[0], file_info[1]]
        res['answer'] = file_info[2]
    return res


def assemble(config_path):
    if not os.path.exists(config_path):
        raise RuntimeError('配置文件不存在')
    # try:
    config = read_config(config_path)
    # except:
    #     raise RuntimeError('配置文件无法解析')

    if not os.path.exists(MAIN_XLS):
        raise RuntimeError('核心表格不存在')
    df = pd.read_excel(MAIN_XLS)
    df[['题目考点']] = df[['题目考点']].astype(str)

    res = [[], [], [], [], int(config['题号起点'])]
    type_arr = [doc_disassembler.QUE_TYPE.GRAMA_SELE.value, doc_disassembler.QUE_TYPE.GRAMA_FISH.value,
                doc_disassembler.QUE_TYPE.GRAMA_ALTER_WORD.value, doc_disassembler.QUE_TYPE.GRAMA_ALTER_SENT.value]

    for m in [0, 2, 3]:
        tmp_type = type_arr[m]
        tmp_df = df[df['题目类型'] == tmp_type]
        has_picked = set()
        for i in range(config['出题数'][tmp_type]):
            point_arr = get_point(config, tmp_type, i)
            tmp_tmp_df = tmp_df.copy().reset_index() if len(point_arr) == 0 else special_select(tmp_df, point_arr)
            tmp_tmp_df = drop_picked(tmp_tmp_df, has_picked)
            tmp_tmp_df = [tmp_tmp_df.iloc[i] for i in range(len(tmp_tmp_df))]
            if len(tmp_tmp_df) == 0:
                raise RuntimeError(tmp_type + '中' + ', '.join(point_arr) + '类型题目不足')
            picked_i = random_arr_with_weight(
                [int(tmp_obj['重要程度']) / (int(tmp_obj['已出次数']) + 1) for tmp_obj in tmp_tmp_df])
            has_picked.add(tmp_tmp_df[picked_i]['题目路径'])
            info_by_file = path_to_question_and_answer(tmp_tmp_df[picked_i]['题目路径'], tmp_type)
            res[m].append({
                'path': tmp_tmp_df[picked_i]['题目路径'],
                'src': [tmp_tmp_df[picked_i]['出处路径'], int(tmp_tmp_df[picked_i]['出处题号'])],
                'point': tmp_tmp_df[picked_i]['题目考点'],
                'importance': int(tmp_tmp_df[picked_i]['重要程度']),
                'times': int(tmp_tmp_df[picked_i]['已出次数']),
                'question': info_by_file['question'],
                'answer': info_by_file['answer']
            })

    tmp_type = type_arr[1]
    tmp_df = df[df['题目类型'] == tmp_type]
    has_picked = set()
    for i in range(config['出题数'][tmp_type]):
        tmp_tmp_df = tmp_df.copy()
        tmp_tmp_df = drop_picked(tmp_tmp_df, has_picked)
        tmp_tmp_df = [tmp_tmp_df.iloc[i] for i in range(len(tmp_tmp_df))]
        if len(tmp_tmp_df) == 0:
            raise RuntimeError(tmp_type + '题目不足')
        picked_i = random_arr_with_weight(
            [1 / (int(tmp_obj['已出次数']) + 1) for tmp_obj in tmp_tmp_df])
        has_picked.add(tmp_tmp_df[picked_i]['题目路径'])
        info_by_file = path_to_question_and_answer(tmp_tmp_df[picked_i]['题目路径'], tmp_type)
        res[1].append({
            'path': tmp_tmp_df[picked_i]['题目路径'],
            'src': [tmp_tmp_df[picked_i]['出处路径'], tmp_tmp_df[picked_i]['出处题号']],
            'times': int(tmp_tmp_df[picked_i]['已出次数']),
            'answer': info_by_file['answer']
        })

    return res


def tab_completion(s, sele_len):
    # return s + '\t' * int(math.ceil((sele_len - len(s)) / TAB_LEN))
    return s + '\t' * 2
    # return s + ' ' * 2 * (sele_len - len(s))


def generate_selection(sele_arr):
    res = ['']
    ok = True
    for i, sele in enumerate(sele_arr):
        if len(sele) > SELE_LEN[4]:
            ok = False
            break
        else:
            res[int(math.floor(i / 4))] += chr(ord('A') + i) + ') ' + tab_completion(sele.strip(), SELE_LEN[4])
    if ok:
        return list(map(lambda item: ' ' * SPACE_NUM + item.strip(), res))

    res = ['', '']
    ok = True
    for i, sele in enumerate(sele_arr):
        if len(sele) > SELE_LEN[2]:
            ok = False
            break
        else:
            res[int(math.floor(i / 2))] += chr(ord('A') + i) + ') ' + tab_completion(sele.strip(), SELE_LEN[2])
    if ok:
        return list(map(lambda item: ' ' * SPACE_NUM + item.strip(), res))

    res = ['', '', '', '']
    for i, sele in enumerate(sele_arr):
        res[i] += chr(ord('A') + i) + ') ' + sele.strip()
    return list(map(lambda item: ' ' * SPACE_NUM + item.strip(), res))


def separate_str(s):
    res = []
    state = 'normal'
    i = 0
    while i < len(s):
        if state == 'normal':
            underline_start = s.find('<_>', i)
            if underline_start == -1:
                res.append((s[i:], 'normal'))
                break
            else:
                res.append((s[i:underline_start], 'normal'))
                i = underline_start + len('<_>')
                state = 'underline'
        elif state == 'underline':
            underline_end = s.find('</_>', i)
            if underline_end == -1:
                raise RuntimeError('下划线标识有始无终')
            else:
                res.append((s[i:underline_end], 'underline'))
                i = underline_end + len('</_>')
                state = 'normal'
    if state != 'normal':
        raise RuntimeError('下划线标识有始无终')
    res = list(filter(lambda item: len(item[0]) > 0, res))
    return res


def str_to_paragraph_inplace(s, para):
    s_arr = separate_str(s)
    for s_item in s_arr:
        if s_item[1] == 'normal':
            para.add_run(s_item[0])
        elif s_item[1] == 'underline':
            para.add_run(s_item[0]).underline = True


def generate_question_inplace(info, type_i, q_num, doc):
    if type_i == 0:
        str_to_paragraph_inplace(str(q_num) + '. ' + re.sub(r'\n', '\n' + ' ' * SPACE_NUM, info[0]),
                                 doc.add_paragraph(''))
        for sele_str in generate_selection(info[1]):
            str_to_paragraph_inplace(sele_str, doc.add_paragraph(''))
    elif type_i == 1:
        doc.add_paragraph(q_num + '. 请至' + q_num + '.docx复制')
    elif type_i == 2:
        str_to_paragraph_inplace(str(q_num) + '. ' + re.sub(r'\n', '\n' + ' ' * SPACE_NUM, info), doc.add_paragraph(''))
    elif type_i == 3:
        str_to_paragraph_inplace(str(q_num) + '. ' + re.sub(r'\n', '\n' + ' ' * SPACE_NUM, info[0]),
                                 doc.add_paragraph(''))
        str_to_paragraph_inplace(' ' * SPACE_NUM + re.sub(r'\n', '\n' + ' ' * SPACE_NUM, info[1]),
                                 doc.add_paragraph(''))


def generate_ans_inplace(info, type_i, q_start, doc):
    if type_i == 0:
        i = 0
        while i < len(info):
            tmp_start = i + q_start
            tmp_end = tmp_start + 4
            tmp_info = ' '.join(info[i:i + 5])
            tmp_str = str(tmp_start) + '-' + str(tmp_end) + '. ' + tmp_info
            str_to_paragraph_inplace(tmp_str, doc.add_paragraph(''))
            i += 5
    elif type_i == 1:
        i = 0
        while i < len(info):
            tmp_start = i + q_start
            tmp_end = tmp_start + 7
            tmp_info = ' '.join(info[i:i + 8])
            tmp_str = str(tmp_start) + '-' + str(tmp_end) + '. ' + tmp_info
            str_to_paragraph_inplace(tmp_str, doc.add_paragraph(''))
            i += 8
    else:
        for i, ans in enumerate(info):
            tmp_str = str(i + q_start) + '. ' + ans
            str_to_paragraph_inplace(tmp_str, doc.add_paragraph(''))


def get_fish_ans_together(info):
    res = []
    for x in info:
        res += x['answer']
    return res


def generate_paper_and_ans(info, dst_dir):
    if not os.path.exists(DST_PATH):
        os.mkdir(DST_PATH)
    dst_dir = DST_PATH + dst_dir + '/'
    if not os.path.exists(dst_dir):
        os.mkdir(dst_dir)
    if not os.path.exists(MAIN_XLS):
        raise RuntimeError('未找到核心表格')
    if not copy_sheet('出卷'):
        raise RuntimeError('备份核心表格失败')

    df = pd.read_excel(MAIN_XLS)

    # 拼卷子
    paper = docx.Document()
    title_arr = ['Choose the best answer (选择最恰当的答案)',
                 'Complete the following passage with the words or phrase in the box. Each can only be used once (将下列单词或词组填入空格。每空格限填一词，每词只能填一次)',
                 'Complete the sentences with the given words in their proper forms (用括号中所给单词的适当形式完成下列句子。每空格限填一词)',
                 'Complete the following sentences as required (根据所给要求完成句子。62－67小题每空格限填一词)']
    q_num = info[4]
    for i in range(4):
        if len(info[i]) == 0:
            continue
        else:
            tmp_title = title_arr[i] + ' (共' + str(len(info[i])) + '分)'
            paper.add_paragraph('').add_run(tmp_title).bold = True
            if i in [0, 2, 3]:
                for j in range(len(info[i])):
                    generate_question_inplace(info[i][j]['question'], i, q_num, paper)
                    q_num += 1
            else:
                for j in range(len(info[i])):
                    tmp_info = str(q_num) + '-' + str(q_num + 7)
                    generate_question_inplace(None, i, tmp_info, paper)
                    shutil.copyfile(info[i][j]['src'][0], dst_dir + tmp_info + '.docx')
                    q_num += 8
    paper.styles['Normal'].font.name = u'Times New Roman'
    paper.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'新宋体')
    for para in paper.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(18)
    paper.save(dst_dir + '试卷.docx')

    # 拼答案
    answer = docx.Document()
    q_num = info[4]
    for i in range(4):
        if len(info[i]) == 0:
            continue
        else:
            if i in [0, 2, 3]:
                generate_ans_inplace([info[i][j]['answer'] for j in range(len(info[i]))], i, q_num, answer)
                q_num += len(info[i])
            else:
                generate_ans_inplace(get_fish_ans_together(info[i]), i, q_num, answer)
                q_num += len(info[i]) * 8
    answer.styles['Normal'].font.name = u'Times New Roman'
    answer.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'新宋体')
    for para in answer.paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Pt(18)
    answer.save(dst_dir + '答案.docx')

    # 拼出处
    src_df = pd.DataFrame(columns=['本卷题号', '题目类型', '出处路径', '出处题号', '题目考点', '重要程度', '已出次数'])
    type_arr = [doc_disassembler.QUE_TYPE.GRAMA_SELE.value, doc_disassembler.QUE_TYPE.GRAMA_FISH.value,
                doc_disassembler.QUE_TYPE.GRAMA_ALTER_WORD.value, doc_disassembler.QUE_TYPE.GRAMA_ALTER_SENT.value]
    q_num = info[4]
    for i in range(4):
        if len(info[i]) == 0:
            continue
        else:
            if i in [0, 2, 3]:
                for j in range(len(info[i])):
                    src_df = src_df.append(pd.Series({
                        '本卷题号': q_num,
                        '题目类型': type_arr[i],
                        '出处路径': info[i][j]['src'][0],
                        '出处题号': info[i][j]['src'][1],
                        '题目考点': info[i][j]['point'],
                        '重要程度': info[i][j]['importance'],
                        '已出次数': info[i][j]['times']
                    }), ignore_index=True)
                    q_num += 1
            else:
                for j in range(len(info[i])):
                    src_df = src_df.append(pd.Series({
                        '本卷题号': str(q_num) + '-' + str(q_num + 7),
                        '题目类型': type_arr[i],
                        '出处路径': info[i][j]['src'][0],
                        '出处题号': info[i][j]['src'][1],
                        '题目考点': '',
                        '重要程度': '',
                        '已出次数': info[i][j]['times']
                    }), ignore_index=True)
                    q_num += 8
    src_df.to_excel(dst_dir + '来源.xlsx', index=False)

    # 改出现次数
    df = pd.read_excel(MAIN_XLS)
    path_arr = []
    for i in range(4):
        for j in range(len(info[i])):
            path_arr.append(info[i][j]['path'])
    for i in range(len(df)):
        if df.iloc[i]['题目路径'] in path_arr:
            df.set_value(i, '已出次数', int(df.iloc[i]['已出次数']) + 1)
    df.to_excel(MAIN_XLS, index=False)
